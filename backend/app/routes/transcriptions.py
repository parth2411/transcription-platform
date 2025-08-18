# backend/app/routes/transcriptions.py - Enhanced with optional titles and export
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, BackgroundTasks, Response
from sqlalchemy.orm import Session
from pydantic import BaseModel, HttpUrl
from typing import List, Optional
import logging
import uuid
import json
import io
from datetime import datetime

from ..database import get_db
from ..models import User, Transcription
from ..services.auth_service import get_current_user
from ..services.transcription_service import TranscriptionService
from ..services.file_service import FileService
from ..config import settings

logger = logging.getLogger(__name__)
router = APIRouter()

# Updated Pydantic models with optional titles
class TranscriptionCreate(BaseModel):
    title: Optional[str] = None  # Made optional
    language: str = "auto"
    generate_summary: bool = True
    speaker_diarization: bool = False
    add_to_knowledge_base: bool = True

class TranscriptionURL(BaseModel):
    url: HttpUrl
    title: Optional[str] = None  # Made optional
    language: str = "auto"
    generate_summary: bool = True
    speaker_diarization: bool = False
    add_to_knowledge_base: bool = True

class TranscriptionText(BaseModel):
    text: str
    title: Optional[str] = None  # Made optional
    generate_summary: bool = True
    add_to_knowledge_base: bool = True

class TranscriptionResponse(BaseModel):
    id: str
    title: str
    status: str
    file_type: Optional[str]
    file_size: Optional[int]
    duration_seconds: Optional[int]
    transcription_text: Optional[str]
    summary_text: Optional[str]
    language: str
    created_at: str
    completed_at: Optional[str]
    processing_time_seconds: Optional[int]
    error_message: Optional[str]

class TranscriptionList(BaseModel):
    transcriptions: List[TranscriptionResponse]
    total: int
    page: int
    per_page: int

# Initialize services
transcription_service = TranscriptionService()
file_service = FileService()

def check_usage_limits(user: User, db: Session):
    """Check if user has exceeded usage limits"""
    if user.subscription_tier == "business":
        return  # Unlimited for business tier
    if user.subscription_tier == "free" and user.monthly_usage >= settings.FREE_TIER_LIMIT:
        raise HTTPException(
            status_code=status.HTTP_402_PAYMENT_REQUIRED,
            detail=f"Monthly limit exceeded. Upgrade to Pro for more transcriptions."
        )
    elif user.subscription_tier == "pro" and user.monthly_usage >= settings.PRO_TIER_LIMIT:
        raise HTTPException(
            status_code=status.HTTP_402_PAYMENT_REQUIRED,
            detail=f"Monthly limit exceeded. Contact support for higher limits."
        )

def increment_usage(user: User, db: Session):
    """Increment user's monthly usage counter"""
    user.monthly_usage += 1
    db.commit()

# Enhanced background processing for routes/transcriptions.py

async def process_transcription_background(
    transcription_id: str,
    processing_type: str,
    file_path_or_url_or_text: str
):
    """Enhanced background task to process transcription with detailed logging"""
    db = None
    try:
        from ..database import SessionLocal
        db = SessionLocal()
        
        # Convert string ID to UUID if needed
        if isinstance(transcription_id, str):
            transcription_id = uuid.UUID(transcription_id)
        
        transcription = db.query(Transcription).filter(
            Transcription.id == transcription_id
        ).first()
        
        if not transcription:
            logger.error(f"Transcription not found: {transcription_id}")
            return
        
        logger.info(f"Starting background processing for {transcription_id}, type: {processing_type}")
        logger.info(f"add_to_knowledge_base flag: {transcription.add_to_knowledge_base}")
        
        # Process based on type
        if processing_type == "file":
            result = await transcription_service.process_file_transcription(
                db, transcription, file_path_or_url_or_text
            )
        elif processing_type == "url":
            result = await transcription_service.process_url_transcription(
                db, transcription, file_path_or_url_or_text
            )
        elif processing_type == "text":
            result = await transcription_service.process_text_transcription(
                db, transcription, file_path_or_url_or_text
            )
        else:
            raise ValueError(f"Unknown processing type: {processing_type}")
        
        # Log final status
        logger.info(f"Background processing completed for {transcription_id}")
        logger.info(f"Final status: {result.status}")
        
        if result.add_to_knowledge_base:
            if result.qdrant_point_ids:
                logger.info(f"Successfully stored {len(result.qdrant_point_ids)} points in knowledge base")
            else:
                logger.warning("Knowledge base storage was enabled but no points were stored")
        
    except Exception as e:
        logger.error(f"Background processing failed for {transcription_id}: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error traceback:", exc_info=True)
        
        # Update transcription status to failed if we have access to it
        if db and 'transcription' in locals():
            try:
                transcription.status = "failed"
                transcription.error_message = str(e)
                db.commit()
                logger.info(f"Updated transcription {transcription_id} status to failed")
            except Exception as update_error:
                logger.error(f"Failed to update transcription status: {update_error}")
    
    finally:
        if db:
            db.close()

@router.post("/upload", response_model=TranscriptionResponse, status_code=status.HTTP_201_CREATED)
async def upload_file_transcription(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),  # Made optional
    language: str = Form("auto"),
    generate_summary: bool = Form(True),
    speaker_diarization: bool = Form(False),
    add_to_knowledge_base: bool = Form(True),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload and process audio/video file for transcription
    """
    try:
        # Check usage limits
        check_usage_limits(current_user, db)
        
        # Validate file
        file_service.validate_file(file.filename, file.size, file.content_type)
        
        # Generate auto title if not provided
        auto_title = title
        if not auto_title or auto_title.strip() == "":
            auto_title = transcription_service.generate_auto_title(file_name=file.filename)
            logger.info(f"Auto-generated title: {auto_title}")
        
        # Upload file to storage
        file_url = await file_service.upload_file(
            file.file, file.filename, file.content_type, str(current_user.id)
        )
        
        # Create transcription record
        transcription = Transcription(
            user_id=current_user.id,
            title=auto_title,
            original_filename=file.filename,
            file_url=file_url,
            file_type=file.content_type,
            file_size=file.size,
            language=language,
            generate_summary=generate_summary,
            speaker_diarization=speaker_diarization,
            add_to_knowledge_base=add_to_knowledge_base,
            status="pending"
        )
        
        db.add(transcription)
        db.commit()
        db.refresh(transcription)
        
        # Increment usage
        increment_usage(current_user, db)
        
        # Download file for processing
        local_file_path = await file_service.download_file(file_url)
        
        # Start background processing
        background_tasks.add_task(
            process_transcription_background,
            str(transcription.id),
            "file",
            local_file_path
        )
        
        logger.info(f"File transcription started: {transcription.id} with title: {transcription.title}")
        
        return TranscriptionResponse(
            id=str(transcription.id),
            title=transcription.title,
            status=transcription.status,
            file_type=transcription.file_type,
            file_size=transcription.file_size,
            duration_seconds=transcription.duration_seconds,
            transcription_text=transcription.transcription_text,
            summary_text=transcription.summary_text,
            language=transcription.language,
            created_at=transcription.created_at.isoformat(),
            completed_at=transcription.completed_at.isoformat() if transcription.completed_at else None,
            processing_time_seconds=transcription.processing_time_seconds,
            error_message=transcription.error_message
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File upload failed: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="File upload failed"
        )

@router.post("/url", response_model=TranscriptionResponse, status_code=status.HTTP_201_CREATED)
async def create_url_transcription(
    background_tasks: BackgroundTasks,
    transcription_data: TranscriptionURL,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Process URL (YouTube, podcast, any video/audio URL) for transcription
    """
    try:
        # Check usage limits
        check_usage_limits(current_user, db)
        
        # Generate auto title if not provided
        auto_title = transcription_data.title
        if not auto_title or auto_title.strip() == "":
            auto_title = transcription_service._generate_fallback_title(str(transcription_data.url))
            logger.info(f"Auto-generated title: {auto_title}")
        
        # Create transcription record
        transcription = Transcription(
            user_id=current_user.id,
            title=auto_title,
            file_url=str(transcription_data.url),
            file_type="url",
            language=transcription_data.language,
            generate_summary=transcription_data.generate_summary,
            speaker_diarization=transcription_data.speaker_diarization,
            add_to_knowledge_base=transcription_data.add_to_knowledge_base,
            status="pending"
        )
        
        db.add(transcription)
        db.commit()
        db.refresh(transcription)
        
        # Increment usage
        increment_usage(current_user, db)
        
        # Start background processing
        background_tasks.add_task(
            process_transcription_background,
            str(transcription.id),
            "url",
            str(transcription_data.url)
        )
        
        logger.info(f"URL transcription started: {transcription.id} with URL: {transcription_data.url}")
        
        return TranscriptionResponse(
            id=str(transcription.id),
            title=transcription.title,
            status=transcription.status,
            file_type=transcription.file_type,
            file_size=transcription.file_size,
            duration_seconds=transcription.duration_seconds,
            transcription_text=transcription.transcription_text,
            summary_text=transcription.summary_text,
            language=transcription.language,
            created_at=transcription.created_at.isoformat(),
            completed_at=transcription.completed_at.isoformat() if transcription.completed_at else None,
            processing_time_seconds=transcription.processing_time_seconds,
            error_message=transcription.error_message
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"URL transcription failed: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="URL transcription failed"
        )

@router.post("/text", response_model=TranscriptionResponse, status_code=status.HTTP_201_CREATED)
async def create_text_transcription(
    background_tasks: BackgroundTasks,
    transcription_data: TranscriptionText,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Process uploaded text for summary and knowledge base
    """
    try:
        # Check usage limits
        check_usage_limits(current_user, db)
        
        # Generate auto title if not provided
        auto_title = transcription_data.title
        if not auto_title or auto_title.strip() == "":
            auto_title = transcription_service.generate_auto_title(transcription_data.text)
            logger.info(f"Auto-generated title: {auto_title}")
        
        # Create transcription record
        transcription = Transcription(
            user_id=current_user.id,
            title=auto_title,
            file_type="text",
            generate_summary=transcription_data.generate_summary,
            add_to_knowledge_base=transcription_data.add_to_knowledge_base,
            status="pending"
        )
        
        db.add(transcription)
        db.commit()
        db.refresh(transcription)
        
        # Increment usage
        increment_usage(current_user, db)
        
        # Start background processing
        background_tasks.add_task(
            process_transcription_background,
            str(transcription.id),
            "text",
            transcription_data.text
        )
        
        logger.info(f"Text processing started: {transcription.id} with title: {transcription.title}")
        
        return TranscriptionResponse(
            id=str(transcription.id),
            title=transcription.title,
            status=transcription.status,
            file_type=transcription.file_type,
            file_size=len(transcription_data.text.encode('utf-8')),
            duration_seconds=0,
            transcription_text=None,  # Will be set after processing
            summary_text=None,  # Will be set after processing
            language="auto",
            created_at=transcription.created_at.isoformat(),
            completed_at=None,
            processing_time_seconds=None,
            error_message=None
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Text processing failed: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Text processing failed"
        )

# Enhanced Export Functionality for All Types
@router.get("/{transcription_id}/export")
async def export_transcription(
    transcription_id: str,
    format: str = "txt",  # txt, json, pdf, srt, docx, csv
    content: str = "both",  # transcription, summary, both
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export transcription and/or summary in various formats
    Supports: txt, json, pdf, srt, docx, csv
    Content options: transcription, summary, both
    """
    try:
        transcription = db.query(Transcription).filter(
            Transcription.id == transcription_id,
            Transcription.user_id == current_user.id
        ).first()
        
        if not transcription:
            raise HTTPException(status_code=404, detail="Transcription not found")
        
        if transcription.status != "completed":
            raise HTTPException(
                status_code=400, 
                detail="Transcription is not completed yet"
            )
        
        # Prepare content based on user selection
        transcription_text = transcription.transcription_text or ""
        summary_text = transcription.summary_text or ""
        
        if content == "transcription":
            main_content = transcription_text
            content_title = "Transcription"
        elif content == "summary":
            main_content = summary_text
            content_title = "Summary"
        else:  # both
            main_content = f"TRANSCRIPTION:\n\n{transcription_text}\n\n" + \
                          f"SUMMARY:\n\n{summary_text}" if summary_text else transcription_text
            content_title = "Transcription & Summary"
        
        # Generate filename
        safe_title = "".join(c for c in transcription.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        base_filename = f"{safe_title}_{content}"
        
        # Handle different export formats
        if format == "txt":
            return _export_as_txt(main_content, base_filename)
        
        elif format == "json":
            return _export_as_json(transcription, content, base_filename)
        
        elif format == "pdf":
            return _export_as_pdf(transcription, main_content, content_title, base_filename)
        
        elif format == "srt":
            if content == "summary":
                raise HTTPException(status_code=400, detail="SRT format only available for transcription")
            return _export_as_srt(transcription_text, base_filename)
        
        elif format == "docx":
            return _export_as_docx(transcription, main_content, content_title, base_filename)
        
        elif format == "csv":
            return _export_as_csv(transcription, base_filename)
        
        else:
            raise HTTPException(
                status_code=400, 
                detail="Unsupported format. Available: txt, json, pdf, srt, docx, csv"
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

def _export_as_txt(content: str, filename: str) -> Response:
    """Export as plain text"""
    return Response(
        content=content,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={filename}.txt"}
    )

def _export_as_json(transcription: Transcription, content_type: str, filename: str) -> Response:
    """Export as JSON with metadata"""
    data = {
        "id": str(transcription.id),
        "title": transcription.title,
        "created_at": transcription.created_at.isoformat(),
        "completed_at": transcription.completed_at.isoformat() if transcription.completed_at else None,
        "duration_seconds": transcription.duration_seconds,
        "language": transcription.language,
        "file_type": transcription.file_type,
        "processing_time_seconds": transcription.processing_time_seconds,
    }
    
    if content_type in ["transcription", "both"]:
        data["transcription"] = transcription.transcription_text
    
    if content_type in ["summary", "both"]:
        data["summary"] = transcription.summary_text
    
    json_content = json.dumps(data, indent=2, ensure_ascii=False)
    
    return Response(
        content=json_content,
        media_type="application/json",
        headers={"Content-Disposition": f"attachment; filename={filename}.json"}
    )

def _export_as_pdf(transcription: Transcription, content: str, content_title: str, filename: str) -> Response:
    """Export as PDF with formatting"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph(transcription.title, title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        metadata = [
            f"<b>Content Type:</b> {content_title}",
            f"<b>Created:</b> {transcription.created_at.strftime('%Y-%m-%d %H:%M')}",
            f"<b>Language:</b> {transcription.language}",
        ]
        
        if transcription.duration_seconds:
            duration_min = transcription.duration_seconds // 60
            duration_sec = transcription.duration_seconds % 60
            metadata.append(f"<b>Duration:</b> {duration_min}:{duration_sec:02d}")
        
        for meta in metadata:
            story.append(Paragraph(meta, styles['Normal']))
        
        story.append(Spacer(1, 20))
        
        # Content
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return Response(
            content=pdf_content,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
        )
        
    except ImportError:
        # Fallback if reportlab not available
        raise HTTPException(
            status_code=500, 
            detail="PDF export requires reportlab package: pip install reportlab"
        )

def _export_as_srt(text: str, filename: str) -> Response:
    """Export as SRT subtitle format"""
    if not text:
        raise HTTPException(status_code=400, detail="No transcription text available")
    
    # Split text into chunks (roughly 10-15 words per subtitle)
    words = text.split()
    chunks = []
    current_chunk = []
    
    for word in words:
        current_chunk.append(word)
        if len(current_chunk) >= 12:  # 12 words per subtitle
            chunks.append(' '.join(current_chunk))
            current_chunk = []
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    # Generate SRT format
    srt_content = ""
    for i, chunk in enumerate(chunks, 1):
        start_time = (i - 1) * 4  # 4 seconds per subtitle
        end_time = i * 4
        
        start_formatted = format_srt_time(start_time)
        end_formatted = format_srt_time(end_time)
        
        srt_content += f"{i}\n"
        srt_content += f"{start_formatted} --> {end_formatted}\n"
        srt_content += f"{chunk}\n\n"
    
    return Response(
        content=srt_content,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={filename}.srt"}
    )

def _export_as_docx(transcription: Transcription, content: str, content_title: str, filename: str) -> Response:
    """Export as DOCX document"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Title
        title = doc.add_heading(transcription.title, 0)
        
        # Metadata table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        metadata_items = [
            ("Content Type", content_title),
            ("Created", transcription.created_at.strftime('%Y-%m-%d %H:%M')),
            ("Language", transcription.language),
        ]
        
        if transcription.duration_seconds:
            duration_min = transcription.duration_seconds // 60
            duration_sec = transcription.duration_seconds % 60
            metadata_items.append(("Duration", f"{duration_min}:{duration_sec:02d}"))
        
        for key, value in metadata_items:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        # Content
        doc.add_heading(content_title, level=1)
        
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
        )
        
    except ImportError:
        raise HTTPException(
            status_code=500, 
            detail="DOCX export requires python-docx package: pip install python-docx"
        )

def _export_as_csv(transcription: Transcription, filename: str) -> Response:
    """Export metadata and content as CSV"""
    import csv
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Headers
    writer.writerow([
        "ID", "Title", "Created", "Completed", "Duration (seconds)", 
        "Language", "File Type", "Status", "Transcription", "Summary"
    ])
    
    # Data
    writer.writerow([
        str(transcription.id),
        transcription.title,
        transcription.created_at.isoformat(),
        transcription.completed_at.isoformat() if transcription.completed_at else "",
        transcription.duration_seconds or "",
        transcription.language,
        transcription.file_type or "",
        transcription.status,
        transcription.transcription_text or "",
        transcription.summary_text or ""
    ])
    
    csv_content = output.getvalue()
    output.close()
    
    return Response(
        content=csv_content,
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}.csv"}
    )

def format_srt_time(seconds: int) -> str:
    """Format seconds into SRT time format (HH:MM:SS,mmm)"""
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:02d},000"

@router.post("/{transcription_id}/share")
async def create_shareable_link(
    transcription_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Create a shareable public link for transcription
    """
    try:
        transcription = db.query(Transcription).filter(
            Transcription.id == transcription_id,
            Transcription.user_id == current_user.id
        ).first()
        
        if not transcription:
            raise HTTPException(status_code=404, detail="Transcription not found")
        
        # Generate a unique share token
        import secrets
        share_token = secrets.token_urlsafe(32)
        
        # Store share token in database (you might want to add a shares table)
        # For now, we'll use a simple approach
        share_url = f"{settings.FRONTEND_URL}/share/{share_token}"
        
        # You could store this in Redis or a shares table
        # For demo purposes, we'll return a basic share URL
        
        return {
            "share_url": share_url,
            "expires_at": "7 days"  # You can implement expiration
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Share link creation failed: {e}")
        raise HTTPException(status_code=500, detail="Share link creation failed")

@router.get("/", response_model=TranscriptionList)
async def list_transcriptions(
    page: int = 1,
    per_page: int = 20,
    status_filter: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    List user's transcriptions with pagination
    """
    try:
        # Build query
        query = db.query(Transcription).filter(Transcription.user_id == current_user.id)
        
        if status_filter:
            query = query.filter(Transcription.status == status_filter)
        
        # Get total count
        total = query.count()
        
        # Apply pagination
        offset = (page - 1) * per_page
        transcriptions = query.order_by(
            Transcription.created_at.desc()
        ).offset(offset).limit(per_page).all()
        
        # Convert to response format
        transcription_list = []
        for t in transcriptions:
            transcription_list.append(TranscriptionResponse(
                id=str(t.id),
                title=t.title,
                status=t.status,
                file_type=t.file_type,
                file_size=t.file_size,
                duration_seconds=t.duration_seconds,
                transcription_text=t.transcription_text,
                summary_text=t.summary_text,
                language=t.language,
                created_at=t.created_at.isoformat(),
                completed_at=t.completed_at.isoformat() if t.completed_at else None,
                processing_time_seconds=t.processing_time_seconds,
                error_message=t.error_message
            ))
        
        return TranscriptionList(
            transcriptions=transcription_list,
            total=total,
            page=page,
            per_page=per_page
        )
        
    except Exception as e:
        logger.error(f"Failed to list transcriptions: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve transcriptions"
        )

@router.get("/{transcription_id}", response_model=TranscriptionResponse)
async def get_transcription(
    transcription_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get specific transcription details
    """
    try:
        transcription = db.query(Transcription).filter(
            Transcription.id == transcription_id,
            Transcription.user_id == current_user.id
        ).first()
        
        if not transcription:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Transcription not found"
            )
        
        return TranscriptionResponse(
            id=str(transcription.id),
            title=transcription.title,
            status=transcription.status,
            file_type=transcription.file_type,
            file_size=transcription.file_size,
            duration_seconds=transcription.duration_seconds,
            transcription_text=transcription.transcription_text,
            summary_text=transcription.summary_text,
            language=transcription.language,
            created_at=transcription.created_at.isoformat(),
            completed_at=transcription.completed_at.isoformat() if transcription.completed_at else None,
            processing_time_seconds=transcription.processing_time_seconds,
            error_message=transcription.error_message
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get transcription {transcription_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve transcription"
        )

@router.delete("/{transcription_id}")
async def delete_transcription(
    transcription_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Delete transcription and associated data
    """
    try:
        transcription = db.query(Transcription).filter(
            Transcription.id == transcription_id,
            Transcription.user_id == current_user.id
        ).first()
        
        if not transcription:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Transcription not found"
            )
        
        # Delete from vector database if exists
        if transcription.qdrant_point_ids:
            await transcription_service.delete_from_qdrant(
                str(current_user.id),
                transcription.qdrant_point_ids
            )
        
        # Delete file from storage if exists
        if transcription.file_url and not transcription.file_url.startswith('http'):
            await file_service.delete_file(transcription.file_url)
        
        # Delete transcription record
        db.delete(transcription)
        db.commit()
        
        logger.info(f"Transcription deleted: {transcription_id}")
        return {"message": "Transcription deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete transcription {transcription_id}: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to delete transcription"
        )

# Realtime transcription endpoint (if you have realtime functionality)
@router.post("/realtime", response_model=TranscriptionResponse, status_code=status.HTTP_201_CREATED)
async def create_realtime_transcription(
    background_tasks: BackgroundTasks,
    audio: UploadFile = File(...),
    title: Optional[str] = Form(None),  # Made optional
    language: str = Form("auto"),
    generate_summary: bool = Form(True),
    add_to_knowledge_base: bool = Form(True),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Process realtime audio recording for transcription
    """
    try:
        # Check usage limits
        check_usage_limits(current_user, db)
        
        # Generate auto title if not provided
        auto_title = title
        if not auto_title or auto_title.strip() == "":
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
            auto_title = f"Realtime Recording - {timestamp}"
            logger.info(f"Auto-generated title: {auto_title}")
        
        # Create transcription record
        transcription = Transcription(
            user_id=current_user.id,
            title=auto_title,
            original_filename=f"realtime_recording_{datetime.now().strftime('%Y%m%d_%H%M%S')}.webm",
            file_type="audio/webm",
            file_size=audio.size,
            language=language,
            generate_summary=generate_summary,
            add_to_knowledge_base=add_to_knowledge_base,
            status="pending"
        )
        
        db.add(transcription)
        db.commit()
        db.refresh(transcription)
        
        # Increment usage
        increment_usage(current_user, db)
        
        # Save audio file temporarily
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as temp_file:
            content = await audio.read()
            temp_file.write(content)
            temp_path = temp_file.name
        
        # Start background processing
        background_tasks.add_task(
            process_transcription_background,
            str(transcription.id),
            "file",
            temp_path
        )
        
        logger.info(f"Realtime transcription started: {transcription.id} with title: {transcription.title}")
        
        return TranscriptionResponse(
            id=str(transcription.id),
            title=transcription.title,
            status=transcription.status,
            file_type=transcription.file_type,
            file_size=transcription.file_size,
            duration_seconds=transcription.duration_seconds,
            transcription_text=transcription.transcription_text,
            summary_text=transcription.summary_text,
            language=transcription.language,
            created_at=transcription.created_at.isoformat(),
            completed_at=transcription.completed_at.isoformat() if transcription.completed_at else None,
            processing_time_seconds=transcription.processing_time_seconds,
            error_message=transcription.error_message
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Realtime transcription failed: {e}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Realtime transcription failed"
        )

# Bulk export endpoint
@router.get("/export/bulk")
async def bulk_export_transcriptions(
    format: str = "json",  # json, csv, zip
    content: str = "both",  # transcription, summary, both
    status_filter: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export multiple transcriptions in bulk
    """
    try:
        # Get user's transcriptions
        query = db.query(Transcription).filter(
            Transcription.user_id == current_user.id,
            Transcription.status == "completed"
        )
        
        if status_filter:
            query = query.filter(Transcription.status == status_filter)
        
        transcriptions = query.order_by(Transcription.created_at.desc()).all()
        
        if not transcriptions:
            raise HTTPException(status_code=404, detail="No completed transcriptions found")
        
        if format == "json":
            return _bulk_export_json(transcriptions, content)
        elif format == "csv":
            return _bulk_export_csv(transcriptions, content)
        elif format == "zip":
            return _bulk_export_zip(transcriptions, content)
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported format. Available: json, csv, zip"
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Bulk export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Bulk export failed: {str(e)}")

def _bulk_export_json(transcriptions: List[Transcription], content_type: str) -> Response:
    """Export all transcriptions as a single JSON file"""
    data = {
        "export_date": datetime.now().isoformat(),
        "total_transcriptions": len(transcriptions),
        "content_type": content_type,
        "transcriptions": []
    }
    
    for t in transcriptions:
        item = {
            "id": str(t.id),
            "title": t.title,
            "created_at": t.created_at.isoformat(),
            "completed_at": t.completed_at.isoformat() if t.completed_at else None,
            "duration_seconds": t.duration_seconds,
            "language": t.language,
            "file_type": t.file_type
        }
        
        if content_type in ["transcription", "both"]:
            item["transcription"] = t.transcription_text
        
        if content_type in ["summary", "both"]:
            item["summary"] = t.summary_text
        
        data["transcriptions"].append(item)
    
    json_content = json.dumps(data, indent=2, ensure_ascii=False)
    
    return Response(
        content=json_content,
        media_type="application/json",
        headers={"Content-Disposition": f"attachment; filename=bulk_export_{content_type}.json"}
    )

def _bulk_export_csv(transcriptions: List[Transcription], content_type: str) -> Response:
    """Export all transcriptions as CSV"""
    import csv
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Headers
    headers = [
        "ID", "Title", "Created", "Completed", "Duration (seconds)", 
        "Language", "File Type", "Status"
    ]
    
    if content_type in ["transcription", "both"]:
        headers.append("Transcription")
    
    if content_type in ["summary", "both"]:
        headers.append("Summary")
    
    writer.writerow(headers)
    
    # Data
    for t in transcriptions:
        row = [
            str(t.id),
            t.title,
            t.created_at.isoformat(),
            t.completed_at.isoformat() if t.completed_at else "",
            t.duration_seconds or "",
            t.language,
            t.file_type or "",
            t.status
        ]
        
        if content_type in ["transcription", "both"]:
            row.append(t.transcription_text or "")
        
        if content_type in ["summary", "both"]:
            row.append(t.summary_text or "")
        
        writer.writerow(row)
    
    csv_content = output.getvalue()
    output.close()
    
    return Response(
        content=csv_content,
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=bulk_export_{content_type}.csv"}
    )

def _bulk_export_zip(transcriptions: List[Transcription], content_type: str) -> Response:
    """Export all transcriptions as individual text files in a ZIP"""
    import zipfile
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for t in transcriptions:
            safe_title = "".join(c for c in t.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            
            if content_type in ["transcription", "both"] and t.transcription_text:
                filename = f"{safe_title}_transcription.txt"
                zip_file.writestr(filename, t.transcription_text)
            
            if content_type in ["summary", "both"] and t.summary_text:
                filename = f"{safe_title}_summary.txt"
                zip_file.writestr(filename, t.summary_text)
            
            # Also include metadata
            metadata = {
                "Title": t.title,
                "Created": t.created_at.isoformat(),
                "Duration": f"{t.duration_seconds} seconds" if t.duration_seconds else "N/A",
                "Language": t.language,
                "File Type": t.file_type or "N/A"
            }
            
            metadata_content = "\n".join([f"{k}: {v}" for k, v in metadata.items()])
            metadata_filename = f"{safe_title}_metadata.txt"
            zip_file.writestr(metadata_filename, metadata_content)
    
    zip_buffer.seek(0)
    
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=bulk_export_{content_type}.zip"}
    )
# Add these debug routes to your existing routes/transcriptions.py file
# Insert these at the end of the file, before any existing routes

@router.get("/debug/qdrant")
async def debug_qdrant_status(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Debug endpoint to check Qdrant connection and status"""
    try:
        from qdrant_client import QdrantClient
        from sentence_transformers import SentenceTransformer
        
        # Test basic connection
        client = QdrantClient(
            url=settings.QDRANT_URL,
            api_key=settings.QDRANT_API_KEY
        )
        
        collection_name = f"user_{current_user.id}_transcriptions"
        
        debug_info = {
            "connection_status": "connected",
            "qdrant_url": settings.QDRANT_URL,
            "target_collection": collection_name,
            "user_id": str(current_user.id)
        }
        
        # Test collections list
        try:
            collections = client.get_collections()
            debug_info["total_collections"] = len(collections.collections)
            debug_info["all_collections"] = [c.name for c in collections.collections]
        except Exception as e:
            debug_info["collections_error"] = str(e)
        
        # Check target collection
        try:
            collection_info = client.get_collection(collection_name)
            debug_info["collection_exists"] = True
            debug_info["points_count"] = collection_info.points_count
            debug_info["vectors_count"] = collection_info.vectors_count
            debug_info["collection_status"] = collection_info.status
        except Exception as e:
            debug_info["collection_exists"] = False
            debug_info["collection_error"] = str(e)
        
        # Test embedder
        try:
            embedder = SentenceTransformer('all-MiniLM-L6-v2')
            test_vector = embedder.encode("test sentence").tolist()
            debug_info["embedder_test"] = {
                "status": "working",
                "vector_size": len(test_vector),
                "sample_values": test_vector[:5]
            }
        except Exception as e:
            debug_info["embedder_test"] = {
                "status": "failed",
                "error": str(e)
            }
        
        # Check database transcriptions
        completed_transcriptions = db.query(Transcription).filter(
            Transcription.user_id == current_user.id,
            Transcription.status == "completed"
        ).count()
        
        kb_transcriptions = db.query(Transcription).filter(
            Transcription.user_id == current_user.id,
            Transcription.status == "completed",
            Transcription.qdrant_point_ids.isnot(None)
        ).count()
        
        debug_info["database_stats"] = {
            "completed_transcriptions": completed_transcriptions,
            "kb_stored_transcriptions": kb_transcriptions,
            "storage_rate": f"{(kb_transcriptions/completed_transcriptions*100):.1f}%" if completed_transcriptions > 0 else "0%"
        }
        
        return debug_info
        
    except Exception as e:
        logger.error(f"Debug endpoint failed: {e}")
        return {
            "connection_status": "failed",
            "error": str(e),
            "error_type": type(e).__name__
        }

@router.post("/debug/test-storage")
async def test_qdrant_storage(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Test storing a sample document in Qdrant"""
    try:
        test_text = "This is a test transcription for debugging knowledge base storage. It contains sample content to verify that the vector embedding and storage system is working correctly."
        test_summary = "Test summary: This debug test verifies that Qdrant storage is functioning properly with vector embeddings."
        
        logger.info("Starting Qdrant storage test...")
        
        point_ids = await transcription_service._store_in_qdrant(
            transcription=test_text,
            summary=test_summary,
            user_id=str(current_user.id),
            transcription_id="debug-test-transcription",
            metadata={
                "title": "Debug Test Transcription",
                "type": "debug_test",
                "created_at": datetime.utcnow().isoformat(),
                "test_mode": True
            }
        )
        
        # Verify storage with a search test
        verification_result = None
        if point_ids:
            try:
                from qdrant_client import QdrantClient
                from sentence_transformers import SentenceTransformer
                
                client = QdrantClient(url=settings.QDRANT_URL, api_key=settings.QDRANT_API_KEY)
                embedder = SentenceTransformer('all-MiniLM-L6-v2')
                
                collection_name = f"user_{current_user.id}_transcriptions"
                test_query_vector = embedder.encode("test debug transcription").tolist()
                
                search_results = client.search(
                    collection_name=collection_name,
                    query_vector=test_query_vector,
                    limit=3
                )
                
                verification_result = {
                    "search_test": "success",
                    "results_found": len(search_results),
                    "top_result_score": search_results[0].score if search_results else 0
                }
                
            except Exception as search_error:
                verification_result = {
                    "search_test": "failed",
                    "error": str(search_error)
                }
        
        return {
            "test_status": "success",
            "points_stored": len(point_ids),
            "point_ids": point_ids,
            "verification": verification_result,
            "message": "Test storage completed successfully"
        }
        
    except Exception as e:
        logger.error(f"Test storage failed: {e}")
        return {
            "test_status": "failed",
            "error": str(e),
            "error_type": type(e).__name__
        }

@router.post("/debug/manual-store/{transcription_id}")
async def manual_store_transcription(
    transcription_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Manually store an existing transcription in Qdrant"""
    try:
        # Get the transcription
        transcription = db.query(Transcription).filter(
            Transcription.id == transcription_id,
            Transcription.user_id == current_user.id,
            Transcription.status == "completed"
        ).first()
        
        if not transcription:
            raise HTTPException(
                status_code=404, 
                detail="Completed transcription not found"
            )
        
        if not transcription.transcription_text:
            raise HTTPException(
                status_code=400,
                detail="Transcription has no text content"
            )
        
        logger.info(f"Manually storing transcription {transcription_id}")
        
        # Store in Qdrant
        point_ids = await transcription_service._store_in_qdrant(
            transcription=transcription.transcription_text,
            summary=transcription.summary_text,
            user_id=str(current_user.id),
            transcription_id=str(transcription.id),
            metadata={
                "title": transcription.title,
                "created_at": transcription.created_at.isoformat(),
                "type": transcription.file_type or "manual_storage",
                "manual_storage": True
            }
        )
        
        result = {
            "status": "success" if point_ids else "failed",
            "points_stored": len(point_ids),
            "point_ids": point_ids,
            "transcription_title": transcription.title
        }
        
        if point_ids:
            # Update the transcription record
            transcription.qdrant_point_ids = point_ids
            transcription.qdrant_collection = f"user_{current_user.id}_transcriptions"
            db.commit()
            result["database_updated"] = True
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Manual store endpoint failed: {e}")
        return {
            "status": "failed",
            "error": str(e),
            "error_type": type(e).__name__
        }

@router.get("/debug/collection-stats")
async def get_collection_statistics(
    current_user: User = Depends(get_current_user)
):
    """Get detailed statistics about the user's Qdrant collection"""
    try:
        from qdrant_client import QdrantClient
        
        client = QdrantClient(
            url=settings.QDRANT_URL,
            api_key=settings.QDRANT_API_KEY
        )
        
        collection_name = f"user_{current_user.id}_transcriptions"
        
        try:
            # Get collection info
            collection_info = client.get_collection(collection_name)
            
            # Get some sample points
            sample_points = client.scroll(
                collection_name=collection_name,
                limit=5,
                with_payload=True,
                with_vectors=False
            )
            
            # Get points by content type
            transcription_points = client.scroll(
                collection_name=collection_name,
                scroll_filter={
                    "must": [
                        {
                            "key": "content_type",
                            "match": {"value": "transcription"}
                        }
                    ]
                },
                limit=1000,
                with_payload=False
            )
            
            summary_points = client.scroll(
                collection_name=collection_name,
                scroll_filter={
                    "must": [
                        {
                            "key": "content_type", 
                            "match": {"value": "summary"}
                        }
                    ]
                },
                limit=1000,
                with_payload=False
            )
            
            return {
                "collection_exists": True,
                "collection_name": collection_name,
                "points_count": collection_info.points_count,
                "vectors_count": collection_info.vectors_count,
                "segments_count": collection_info.segments_count,
                "status": collection_info.status,
                "transcription_points": len(transcription_points[0]),
                "summary_points": len(summary_points[0]),
                "sample_points": [
                    {
                        "id": str(point.id),
                        "payload_keys": list(point.payload.keys()) if point.payload else [],
                        "content_type": point.payload.get("content_type") if point.payload else None,
                        "title": point.payload.get("title") if point.payload else None,
                        "text_preview": point.payload.get("text_preview", "")[:100] if point.payload else ""
                    }
                    for point in sample_points[0][:5]
                ]
            }
            
        except Exception as e:
            return {
                "collection_exists": False,
                "collection_name": collection_name,
                "error": str(e),
                "error_type": type(e).__name__
            }
        
    except Exception as e:
        logger.error(f"Collection stats endpoint failed: {e}")
        return {"error": str(e)}

@router.post("/debug/fix-existing")
async def fix_existing_transcriptions(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Fix existing transcriptions that weren't stored in knowledge base"""
    try:
        # Find completed transcriptions without knowledge base storage
        missing_kb_transcriptions = db.query(Transcription).filter(
            Transcription.user_id == current_user.id,
            Transcription.status == "completed",
            Transcription.transcription_text.isnot(None),
            Transcription.qdrant_point_ids.is_(None)
        ).all()
        
        if not missing_kb_transcriptions:
            return {
                "status": "no_work_needed",
                "message": "All completed transcriptions are already in knowledge base",
                "processed": 0
            }
        
        results = {
            "status": "processing",
            "total_found": len(missing_kb_transcriptions),
            "processed": 0,
            "successful": 0,
            "failed": 0,
            "details": []
        }
        
        for transcription in missing_kb_transcriptions:
            try:
                logger.info(f"Processing transcription {transcription.id}: {transcription.title}")
                
                point_ids = await transcription_service._store_in_qdrant(
                    transcription=transcription.transcription_text,
                    summary=transcription.summary_text,
                    user_id=str(current_user.id),
                    transcription_id=str(transcription.id),
                    metadata={
                        "title": transcription.title,
                        "created_at": transcription.created_at.isoformat(),
                        "type": transcription.file_type or "retroactive_storage",
                        "retroactive_fix": True
                    }
                )
                
                if point_ids:
                    # Update the transcription record
                    transcription.qdrant_point_ids = point_ids
                    transcription.qdrant_collection = f"user_{current_user.id}_transcriptions"
                    db.commit()
                    
                    results["successful"] += 1
                    results["details"].append({
                        "id": str(transcription.id),
                        "title": transcription.title,
                        "status": "success",
                        "points_stored": len(point_ids)
                    })
                else:
                    results["failed"] += 1
                    results["details"].append({
                        "id": str(transcription.id), 
                        "title": transcription.title,
                        "status": "failed",
                        "error": "No points returned from storage"
                    })
                
                results["processed"] += 1
                
            except Exception as e:
                results["failed"] += 1
                results["details"].append({
                    "id": str(transcription.id),
                    "title": transcription.title, 
                    "status": "failed",
                    "error": str(e)
                })
                logger.error(f"Failed to process transcription {transcription.id}: {e}")
        
        results["status"] = "completed"
        return results
        
    except Exception as e:
        logger.error(f"Fix existing transcriptions failed: {e}")
        return {
            "status": "failed",
            "error": str(e),
            "error_type": type(e).__name__
        }